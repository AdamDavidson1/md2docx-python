
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_test_document(output_path):
    """Create a test Word document with various formatting"""
    doc = Document()

    # Add a title
    doc.add_heading("Test Document", level=1)

    # Add some regular paragraphs
    doc.add_paragraph("This is a regular paragraph with some text.")

    # Add formatted text
    p = doc.add_paragraph()
    p.add_run("This paragraph has ")
    p.add_run("bold text, ").bold = True
    p.add_run("italic text, ").italic = True
    run = p.add_run("and bold-italic text.")
    run.bold = True
    run.italic = True

    # Add different heading levels
    doc.add_heading("Heading Level 2", level=2)
    doc.add_heading("Heading Level 3", level=3)

    # Add bullet points
    doc.add_paragraph("First bullet point", style="List Bullet")
    doc.add_paragraph("Second bullet point", style="List Bullet")

    # Add numbered list
    doc.add_paragraph("First numbered item", style="List Number")
    doc.add_paragraph("Second numbered item", style="List Number")

    # Add a code block
    code_block = doc.add_paragraph()
    code_block.style = doc.styles["Normal"]
    code_run = code_block.add_run(
        """def hello_world():
    print("Hello, World!")
    return True"""
    )
    code_run.font.name = "Consolas"

    # Add inline code
    p = doc.add_paragraph("Here is some inline code: ")
    code = p.add_run('print("Hello")')
    code.font.name = "Courier New"

    # Save the document
    doc.save(output_path)


if __name__ == "__main__":
    create_test_document("sample_files/test_document.docx")
